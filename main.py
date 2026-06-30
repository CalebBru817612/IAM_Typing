import os
import threading
from datetime import datetime
import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk

try:
    import vlc
except Exception:
    vlc = None

try:
    import whisper
except Exception:
    whisper = None


APP_NAME = "IAM Typing"
APP_VERSION = "v1.0.0-compact"

COLOR_BG = "#eef2f7"
COLOR_PANEL = "#f8fafc"
COLOR_PANEL_DARK = "#ffffff"
COLOR_BORDER = "#d8dde6"
COLOR_SIDEBAR = "#071f3d"
COLOR_SIDEBAR_BTN = "#123a6f"
COLOR_SIDEBAR_BTN_HOVER = "#1d4ed8"
COLOR_ACCENT = "#f59e0b"
COLOR_BLUE = COLOR_SIDEBAR_BTN
COLOR_BLUE_2 = "#0b2347"
COLOR_WHITE = "#ffffff"
COLOR_TEXT = "#0b2347"
COLOR_MUTED = "#64748b"
COLOR_DANGER = "#b91c1c"
COLOR_ENTRY = "#ffffff"


SUPPORTED_FILETYPES = [
    ("Supported Media Files", "*.mp3 *.wav *.m4a *.aac *.flac *.opus *.ogg *.wma *.mp4 *.mov *.mkv *.avi *.webm *.mpeg *.mpg *.3gp *.DS2 *.ds2"),
    ("Audio Files", "*.mp3 *.wav *.m4a *.aac *.flac *.opus *.ogg *.wma *.DS2 *.ds2"),
    ("Video Files", "*.mp4 *.mov *.mkv *.avi *.webm *.mpeg *.mpg *.3gp"),
    ("All Files", "*.*"),
]


class IAMTypingApp:
    def __init__(self, root):
        self.root = root
        self.root.title(f"{APP_NAME} {APP_VERSION}")
        self.root.geometry("1280x760")
        self.root.minsize(980, 600)
        self.root.configure(bg=COLOR_BG)

        self.selected_file = ""
        self.recent_files = {}
        self.is_transcribing = False
        self.whisper_model = None

        if vlc is not None:
            self.player = vlc.MediaPlayer()
        else:
            self.player = None

        self.setup_style()
        self.build_layout()
        self.update_timeline()

    def setup_style(self):
        style = ttk.Style()
        style.theme_use("default")
        style.configure(
            "blue.Horizontal.TProgressbar",
            troughcolor=COLOR_PANEL_DARK,
            background=COLOR_ACCENT,
            bordercolor=COLOR_PANEL_DARK,
            lightcolor=COLOR_ACCENT,
            darkcolor=COLOR_ACCENT,
            thickness=12
        )

    def build_layout(self):
        self.root.grid_columnconfigure(1, weight=1)
        self.root.grid_rowconfigure(0, weight=1)

        self.sidebar = tk.Frame(self.root, bg=COLOR_SIDEBAR, width=220)
        self.sidebar.grid(row=0, column=0, sticky="ns")
        self.sidebar.grid_propagate(False)

        self.main_area = tk.Frame(self.root, bg=COLOR_BG)
        self.main_area.grid(row=0, column=1, sticky="nsew")
        self.main_area.grid_columnconfigure(0, weight=1)
        self.main_area.grid_rowconfigure(1, weight=1)

        self.build_sidebar()
        self.build_topbar()
        self.build_content()

    def build_sidebar(self):
        logo_frame = tk.Frame(self.sidebar, bg=COLOR_SIDEBAR)
        logo_frame.pack(fill="x", pady=(20, 14))

        tk.Label(
            logo_frame,
            text="IAM Typing",
            bg=COLOR_SIDEBAR,
            fg=COLOR_WHITE,
            font=("Arial", 20, "bold")
        ).pack(anchor="w", padx=25)

        tk.Label(
            logo_frame,
            text="Single-speaker typing workspace",
            bg=COLOR_SIDEBAR,
            fg="#9fb3cc",
            font=("Arial", 9)
        ).pack(anchor="w", padx=26, pady=(4, 0))

        tk.Label(
            logo_frame,
            text=APP_VERSION,
            bg=COLOR_SIDEBAR,
            fg=COLOR_ACCENT,
            font=("Arial", 9, "bold")
        ).pack(anchor="w", padx=26, pady=(3, 0))

        tk.Frame(self.sidebar, bg=COLOR_ACCENT, height=2).pack(fill="x", padx=20, pady=(10, 14))

        sidebar_buttons = [
            ("Start Fresh", self.start_fresh),
            ("Files / History", self.show_history),
            ("Playback", self.focus_playback),
            ("Export", self.focus_export),
            ("Settings", self.show_settings),
            ("About", self.show_about),
        ]

        for text, command in sidebar_buttons:
            btn = tk.Button(
                self.sidebar,
                text=text,
                font=("Arial", 12, "bold"),
                bg=COLOR_SIDEBAR_BTN,
                fg=COLOR_WHITE,
                activebackground=COLOR_SIDEBAR_BTN_HOVER,
                activeforeground=COLOR_WHITE,
                relief="flat",
                bd=0,
                width=19,
                height=1,
                cursor="hand2",
                command=command
            )
            btn.pack(pady=4, padx=12)

        tk.Label(
            self.sidebar,
            text="Single-speaker dictation\nNo speaker detection",
            font=("Arial", 9),
            fg="#9fb3cc",
            bg=COLOR_SIDEBAR,
            justify="left"
        ).pack(side="bottom", anchor="w", padx=22, pady=14)

    def build_topbar(self):
        top_bar = tk.Frame(
            self.main_area,
            bg=COLOR_WHITE,
            height=66,
            highlightbackground=COLOR_BORDER,
            highlightthickness=1
        )
        top_bar.grid(row=0, column=0, columnspan=2, sticky="ew")
        top_bar.grid_propagate(False)
        top_bar.grid_columnconfigure(0, weight=1)

        title_wrap = tk.Frame(top_bar, bg=COLOR_WHITE)
        title_wrap.grid(row=0, column=0, sticky="w", padx=22, pady=(8, 0))

        tk.Label(
            title_wrap,
            text="IAM Typing",
            font=("Arial", 21, "bold"),
            fg=COLOR_TEXT,
            bg=COLOR_WHITE
        ).pack(anchor="w")

        tk.Label(
            title_wrap,
            text="Single-speaker dictation workspace",
            font=("Arial", 10),
            fg=COLOR_MUTED,
            bg=COLOR_WHITE
        ).pack(anchor="w", pady=(2, 0))

        self.top_status = tk.Label(
            top_bar,
            text="Ready",
            font=("Arial", 10, "bold"),
            fg=COLOR_WHITE,
            bg=COLOR_SIDEBAR_BTN,
            padx=16,
            pady=8
        )
        self.top_status.grid(row=0, column=1, sticky="e", padx=22, pady=(14, 0))

    def build_content(self):
        # The outer page scrollbar keeps the full workspace usable on smaller screens.
        self.main_area.grid_columnconfigure(0, weight=1)
        self.main_area.grid_columnconfigure(1, weight=0)
        self.main_area.grid_rowconfigure(1, weight=1)

        page_canvas = tk.Canvas(
            self.main_area,
            bg=COLOR_BG,
            highlightthickness=0
        )
        page_scrollbar = ttk.Scrollbar(
            self.main_area,
            orient="vertical",
            command=page_canvas.yview
        )

        page_canvas.grid(row=1, column=0, sticky="nsew", padx=(14, 0), pady=14)
        page_scrollbar.grid(row=1, column=1, sticky="ns", padx=(5, 10), pady=14)

        content = tk.Frame(page_canvas, bg=COLOR_BG)
        canvas_window = page_canvas.create_window((0, 0), window=content, anchor="nw")

        def update_scroll_region(event=None):
            page_canvas.configure(scrollregion=page_canvas.bbox("all"))

        def update_canvas_width(event):
            page_canvas.itemconfig(canvas_window, width=event.width)

        def mouse_wheel_scroll(event):
            page_canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

        def bind_mouse_wheel(event):
            page_canvas.bind_all("<MouseWheel>", mouse_wheel_scroll)

        def unbind_mouse_wheel(event):
            page_canvas.unbind_all("<MouseWheel>")

        content.bind("<Configure>", update_scroll_region)
        content.bind("<Enter>", bind_mouse_wheel)
        content.bind("<Leave>", unbind_mouse_wheel)
        page_canvas.bind("<Configure>", update_canvas_width)
        page_canvas.configure(yscrollcommand=page_scrollbar.set)

        content.grid_columnconfigure(0, weight=1)
        content.grid_rowconfigure(2, weight=1)

        self.content = content
        self.page_canvas = page_canvas

        self.build_file_panels(content)
        self.build_progress_panel(content)
        self.build_work_area(content)
        self.build_export_panel(content)

    def make_panel(self, parent, height=None):
        panel = tk.Frame(
            parent,
            bg=COLOR_PANEL,
            highlightbackground=COLOR_BORDER,
            highlightthickness=1
        )
        if height:
            panel.configure(height=height)
            panel.grid_propagate(False)
            panel.pack_propagate(False)
        return panel

    def build_file_panels(self, parent):
        top_panels = tk.Frame(parent, bg=COLOR_BG)
        top_panels.grid(row=0, column=0, sticky="ew")
        top_panels.grid_columnconfigure(0, weight=1)
        top_panels.grid_columnconfigure(1, weight=1)
        top_panels.grid_columnconfigure(2, weight=1)

        upload_panel = self.make_panel(top_panels, 150)
        upload_panel.grid(row=0, column=0, sticky="ew", padx=(0, 10))

        tk.Label(
            upload_panel,
            text="Select Dictation File",
            font=("Segoe UI Semibold", 13),
            fg=COLOR_TEXT,
            bg=COLOR_PANEL
        ).pack(anchor="w", padx=16, pady=(12, 2))

        tk.Label(
            upload_panel,
            text="Load one audio/video dictation file. Then use playback or Auto Draft.",
            font=("Segoe UI", 9),
            fg=COLOR_MUTED,
            bg=COLOR_PANEL,
            wraplength=310,
            justify="left"
        ).pack(anchor="w", padx=16, pady=(0, 8))

        tk.Button(
            upload_panel,
            text="Browse Media File",
            font=("Segoe UI Semibold", 10),
            bg=COLOR_BLUE,
            fg=COLOR_WHITE,
            activebackground=COLOR_SIDEBAR_BTN_HOVER,
            activeforeground=COLOR_WHITE,
            relief="flat",
            bd=0,
            padx=18,
            pady=7,
            cursor="hand2",
            command=self.select_file
         ).pack(anchor="w", padx=16)

        tk.Label(
            upload_panel,
            text="MP3 • WAV • M4A • OPUS • MP4 • MKV • DS2 listed",
            font=("Segoe UI", 9),
            fg=COLOR_MUTED,
            bg=COLOR_PANEL
        ).pack(anchor="w", padx=16, pady=(8, 0))

        info_panel = self.make_panel(top_panels, 150)
        info_panel.grid(row=0, column=1, sticky="ew", padx=(0, 10))

        tk.Label(
            info_panel,
            text="File Information",
            font=("Segoe UI Semibold", 11),
            fg=COLOR_BLUE_2,
            bg=COLOR_PANEL
        ).pack(anchor="w", padx=16, pady=(10, 6))

        self.file_name_value = tk.Label(
            info_panel,
            text="Name: No file selected",
            font=("Segoe UI", 9),
            fg=COLOR_TEXT,
            bg=COLOR_PANEL,
            wraplength=340,
            justify="left"
        )
        self.file_name_value.pack(anchor="w", padx=16, pady=3)

        self.file_size_value = tk.Label(
            info_panel,
            text="Size: -- MB",
            font=("Segoe UI", 9),
            fg=COLOR_TEXT,
            bg=COLOR_PANEL
        )
        self.file_size_value.pack(anchor="w", padx=16, pady=3)

        self.file_format_value = tk.Label(
            info_panel,
            text="Format: ---",
            font=("Segoe UI", 9),
            fg=COLOR_TEXT,
            bg=COLOR_PANEL
        )
        self.file_format_value.pack(anchor="w", padx=16, pady=3)

        settings_panel = self.make_panel(top_panels, 150)
        settings_panel.grid(row=0, column=2, sticky="ew")

        tk.Label(
            settings_panel,
            text="Dictation Mode",
            font=("Segoe UI Semibold", 11),
            fg=COLOR_BLUE_2,
            bg=COLOR_PANEL
        ).pack(anchor="w", padx=16, pady=(10, 6))

        for item in ["Whisper Base", "Single Speaker", "Typing Proofread Workflow"]:
            tk.Label(
                settings_panel,
                text=item,
                font=("Arial", 9),
                fg=COLOR_TEXT,
                bg=COLOR_WHITE,
                padx=10,
                pady=5,
                highlightbackground=COLOR_BORDER,
                highlightthickness=1
            ).pack(fill="x", padx=16, pady=2)

        self.start_button = tk.Button(
            settings_panel,
            text="Start Auto Draft",
            font=("Segoe UI", 10, "bold"),
            bg=COLOR_BLUE,
            fg=COLOR_WHITE,
            activebackground=COLOR_SIDEBAR_BTN_HOVER,
            activeforeground=COLOR_WHITE,
            relief="flat",
            bd=0,
            padx=11,
            pady=5,
            cursor="hand2",
            command=self.start_auto_draft
        )
        self.start_button.pack(fill="x", padx=55, pady=(6, 0))

    def build_progress_panel(self, parent):
        progress_panel = self.make_panel(parent, 66)
        progress_panel.grid(row=1, column=0, sticky="ew", pady=10)

        self.progress_title = tk.Label(
            progress_panel,
            text="Ready",
            font=("Segoe UI Semibold", 11),
            fg=COLOR_BLUE_2,
            bg=COLOR_PANEL
        )
        self.progress_title.pack(anchor="w", padx=16, pady=(8, 3))

        self.progress = ttk.Progressbar(
            progress_panel,
            orient="horizontal",
            mode="determinate",
            style="blue.Horizontal.TProgressbar"
        )
        self.progress.pack(fill="x", padx=16)

        self.progress_percent = tk.Label(
            progress_panel,
            text="0%",
            font=("Segoe UI", 9),
            fg=COLOR_MUTED,
            bg=COLOR_PANEL
        )
        self.progress_percent.pack(anchor="e", padx=16, pady=(2, 0))

    def build_work_area(self, parent):
        work_area = tk.Frame(parent, bg=COLOR_BG, height=430)
        work_area.grid(row=2, column=0, sticky="nsew")
        work_area.grid_propagate(False)
        work_area.grid_columnconfigure(0, weight=3)
        work_area.grid_columnconfigure(1, weight=2)
        work_area.grid_rowconfigure(0, weight=1)

        transcript_panel = self.make_panel(work_area)
        transcript_panel.grid(row=0, column=0, sticky="nsew", padx=(0, 10))
        transcript_panel.grid_columnconfigure(0, weight=1)
        transcript_panel.grid_rowconfigure(2, weight=1)

        header = tk.Frame(transcript_panel, bg=COLOR_PANEL)
        header.grid(row=0, column=0, sticky="ew", padx=16, pady=(10, 4))
        header.grid_columnconfigure(0, weight=1)

        tk.Label(
            header,
            text="Typed Document / Auto Draft",
            font=("Segoe UI Semibold", 13),
            fg=COLOR_BLUE_2,
            bg=COLOR_PANEL
        ).grid(row=0, column=0, sticky="w")

        self.word_count_label = tk.Label(
            header,
            text="Words: 0",
            font=("Segoe UI", 9),
            fg=COLOR_MUTED,
            bg=COLOR_PANEL
        )
        self.word_count_label.grid(row=0, column=1, sticky="e")

        self.search_box = tk.Entry(
            transcript_panel,
            bg=COLOR_ENTRY,
            fg=COLOR_TEXT,
            insertbackground=COLOR_TEXT,
            relief="flat",
            font=("Segoe UI", 10)
        )
        self.search_box.grid(row=1, column=0, sticky="ew", padx=16, pady=(0, 6), ipady=4)
        self.search_box.insert(0, "Search text here...")

        text_wrap = tk.Frame(transcript_panel, bg=COLOR_PANEL)
        text_wrap.grid(row=2, column=0, sticky="nsew", padx=16, pady=(0, 8))
        text_wrap.grid_columnconfigure(0, weight=1)
        text_wrap.grid_rowconfigure(0, weight=1)

        self.transcript_text = tk.Text(
            text_wrap,
            bg=COLOR_PANEL_DARK,
            fg=COLOR_TEXT,
            insertbackground=COLOR_TEXT,
            relief="flat",
            font=("Consolas", 10),
            wrap="word",
            padx=12,
            pady=12,
            spacing1=2,
            spacing2=4,
            spacing3=2
        )
        self.transcript_text.grid(row=0, column=0, sticky="nsew")

        text_scroll = tk.Scrollbar(text_wrap, command=self.transcript_text.yview)
        text_scroll.grid(row=0, column=1, sticky="ns")
        self.transcript_text.config(yscrollcommand=text_scroll.set)

        self.transcript_text.insert("end", "Your typed document or auto draft will appear here...")
        self.transcript_text.bind("<KeyRelease>", self.update_word_count)

        action_bar = tk.Frame(transcript_panel, bg=COLOR_PANEL)
        action_bar.grid(row=3, column=0, sticky="ew", padx=16, pady=(0, 10))

        tk.Button(
            action_bar,
            text="Copy Text",
            font=("Segoe UI", 9, "bold"),
            bg=COLOR_BLUE,
            fg=COLOR_WHITE,
            activebackground=COLOR_SIDEBAR_BTN_HOVER,
            activeforeground=COLOR_WHITE,
            relief="flat",
            bd=0,
            padx=11,
            pady=5,
            cursor="hand2",
            command=self.copy_text
        ).pack(side="left")

        tk.Button(
            action_bar,
            text="Clear Text",
            font=("Segoe UI", 9, "bold"),
            bg=COLOR_SIDEBAR_BTN,
            fg=COLOR_WHITE,
            activebackground=COLOR_DANGER,
            activeforeground=COLOR_WHITE,
            relief="flat",
            bd=0,
            padx=11,
            pady=5,
            cursor="hand2",
            command=self.clear_text
        ).pack(side="left", padx=8)

        self.playback_panel = self.make_panel(work_area)
        self.playback_panel.grid(row=0, column=1, sticky="nsew")
        self.build_playback_panel(self.playback_panel)

    def build_playback_panel(self, panel):
        panel.grid_columnconfigure(0, weight=1)

        tk.Label(
            panel,
            text="Playback",
            font=("Segoe UI Semibold", 13),
            fg=COLOR_BLUE_2,
            bg=COLOR_PANEL
        ).grid(row=0, column=0, sticky="w", padx=16, pady=(10, 5))

        button_row = tk.Frame(panel, bg=COLOR_PANEL)
        button_row.grid(row=1, column=0, pady=(4, 6))

        buttons = [
            ("⏪ 5s", self.rewind_audio),
            ("▶", self.play_audio),
            ("⏸", self.pause_audio),
            ("⏹", self.stop_audio),
            ("5s ⏩", self.forward_audio),
        ]

        for text, command in buttons:
            tk.Button(
                button_row,
                text=text,
                font=("Segoe UI Symbol", 12, "bold"),
                bg=COLOR_ACCENT if text == "▶" else COLOR_SIDEBAR_BTN,
                fg=COLOR_WHITE,
                activebackground=COLOR_SIDEBAR_BTN_HOVER,
                activeforeground=COLOR_WHITE,
                relief="flat",
                bd=0,
                width=6,
                height=1,
                cursor="hand2",
                command=command
            ).pack(side="left", padx=3)

        self.timeline = tk.Scale(
            panel,
            from_=0,
            to=100,
            orient="horizontal",
            showvalue=False,
            bg=COLOR_PANEL,
            fg=COLOR_WHITE,
            troughcolor=COLOR_ACCENT,
            highlightthickness=0,
            borderwidth=0,
            length=310
        )
        self.timeline.grid(row=2, column=0, sticky="ew", padx=22, pady=(4, 1))
        self.timeline.bind("<ButtonRelease-1>", lambda event: self.seek_audio(self.timeline.get()))

        time_frame = tk.Frame(panel, bg=COLOR_PANEL)
        time_frame.grid(row=3, column=0, sticky="ew", padx=24)
        time_frame.grid_columnconfigure(1, weight=1)

        self.current_time = tk.Label(time_frame, text="00:00", font=("Segoe UI", 9), fg=COLOR_MUTED, bg=COLOR_PANEL)
        self.current_time.grid(row=0, column=0, sticky="w")

        self.total_time = tk.Label(time_frame, text="00:00", font=("Segoe UI", 9), fg=COLOR_MUTED, bg=COLOR_PANEL)
        self.total_time.grid(row=0, column=2, sticky="e")

        audio_controls = tk.Frame(panel, bg=COLOR_PANEL)
        audio_controls.grid(row=4, column=0, sticky="ew", padx=24, pady=(12, 0))
        audio_controls.grid_columnconfigure(1, weight=1)

        tk.Label(audio_controls, text="Volume", font=("Segoe UI", 9), fg=COLOR_MUTED, bg=COLOR_PANEL).grid(row=0, column=0, sticky="w")

        self.volume_slider = tk.Scale(
            audio_controls,
            from_=0,
            to=100,
            orient="horizontal",
            showvalue=False,
            bg=COLOR_PANEL,
            troughcolor=COLOR_ACCENT,
            highlightthickness=0,
            borderwidth=0,
            length=130,
            command=lambda value: self.set_volume(value)
        )
        self.volume_slider.set(80)
        self.volume_slider.grid(row=0, column=1, sticky="ew", padx=10)

        tk.Label(audio_controls, text="Speed", font=("Segoe UI", 9), fg=COLOR_MUTED, bg=COLOR_PANEL).grid(row=1, column=0, sticky="w", pady=(8, 0))

        self.speed_menu = tk.StringVar(value="1.0x")
        speed_dropdown = tk.OptionMenu(
            audio_controls,
            self.speed_menu,
            "0.75x",
            "1.0x",
            "1.25x",
            "1.5x",
            "2.0x",
            command=self.change_speed
        )
        speed_dropdown.config(
            bg=COLOR_SIDEBAR_BTN,
            fg=COLOR_WHITE,
            activebackground=COLOR_SIDEBAR_BTN_HOVER,
            activeforeground=COLOR_WHITE,
            borderwidth=0,
            highlightthickness=0
        )
        speed_dropdown.grid(row=1, column=1, sticky="w", padx=10, pady=(8, 0))

        note = tk.Label(
            panel,
            text="Tip: Use playback for manual typing, or Auto Draft first and proofread.",
            font=("Segoe UI", 10),
            fg=COLOR_MUTED,
            bg=COLOR_PANEL,
            wraplength=330,
            justify="left"
        )
        note.grid(row=5, column=0, sticky="w", padx=24, pady=(16, 0))

    def build_export_panel(self, parent):
        export_panel = self.make_panel(parent, 82)
        export_panel.grid(row=3, column=0, sticky="ew", pady=(10, 0))

        tk.Label(
            export_panel,
            text="Export Document",
            font=("Segoe UI Semibold", 11),
            fg=COLOR_BLUE_2,
            bg=COLOR_PANEL
        ).pack(anchor="w", padx=16, pady=(9, 4))

        export_buttons_frame = tk.Frame(export_panel, bg=COLOR_PANEL)
        export_buttons_frame.pack(fill="x", padx=16, pady=(0, 10))

        for text, command in [
            ("Export as DOCX", self.export_docx),
            ("Export as TXT", self.export_txt),
        ]:
            tk.Button(
                export_buttons_frame,
                text=text,
                font=("Segoe UI", 10),
                bg=COLOR_SIDEBAR_BTN,
                fg=COLOR_WHITE,
                activebackground=COLOR_SIDEBAR_BTN_HOVER,
                activeforeground=COLOR_WHITE,
                relief="flat",
                bd=0,
                padx=15,
                pady=7,
                cursor="hand2",
                command=command
            ).pack(side="left", padx=(0, 10))

    def select_file(self):
        selected_file = filedialog.askopenfilename(
            title="Select Dictation Audio or Video File",
            filetypes=SUPPORTED_FILETYPES
        )

        if not selected_file:
            return

        self.selected_file = selected_file
        self.load_media_for_playback(selected_file)
        self.update_file_info(selected_file)

        file_name = os.path.basename(selected_file)
        self.recent_files[file_name] = selected_file

        if len(self.recent_files) > 15:
            oldest = list(self.recent_files.keys())[0]
            del self.recent_files[oldest]

        self.set_status("File ready")
        self.progress_title.config(text="File ready. Use playback or Start Auto Draft.")

    def update_file_info(self, file_path):
        file_name = os.path.basename(file_path)
        size_mb = round(os.path.getsize(file_path) / (1024 * 1024), 2)
        extension = os.path.splitext(file_path)[1].upper().replace(".", "")

        self.file_name_value.config(text=f"Name: {file_name}")
        self.file_size_value.config(text=f"Size: {size_mb} MB")
        self.file_format_value.config(text=f"Format: {extension}")

    def load_media_for_playback(self, file_path):
        if self.player is None or vlc is None:
            return

        try:
            media = vlc.Media(file_path)
            self.player.set_media(media)
            self.player.audio_set_volume(int(self.volume_slider.get()))
        except Exception as error:
            messagebox.showerror("Playback Error", f"Could not load audio:\n\n{error}")

    def play_audio(self):
        if not self.selected_file:
            messagebox.showerror("No file", "Please select a file first.")
            return

        if self.player is None:
            messagebox.showerror("VLC missing", "python-vlc/VLC is not available.")
            return

        if self.player.get_media() is None:
            self.load_media_for_playback(self.selected_file)

        self.player.play()

    def pause_audio(self):
        if self.player:
            self.player.pause()

    def stop_audio(self):
        if self.player:
            self.player.stop()
            self.timeline.set(0)
            self.current_time.config(text="00:00")

    def rewind_audio(self):
        self.jump_audio(-5000)

    def forward_audio(self):
        self.jump_audio(5000)

    def jump_audio(self, milliseconds):
        if not self.player:
            return

        current = self.player.get_time()
        length = self.player.get_length()

        if current < 0:
            current = 0

        new_time = current + milliseconds

        if new_time < 0:
            new_time = 0

        if length > 0 and new_time > length:
            new_time = length

        self.player.set_time(new_time)

    def seek_audio(self, value):
        if self.player:
            self.player.set_position(float(value) / 100)

    def set_volume(self, value):
        if self.player:
            self.player.audio_set_volume(int(float(value)))

    def change_speed(self, selection):
        if not self.player:
            return

        speed = float(selection.replace("x", ""))
        self.player.set_rate(speed)

    def update_timeline(self):
        if self.player is not None:
            try:
                length = self.player.get_length()
                current = self.player.get_time()

                if length > 0 and current >= 0:
                    timeline_position = (current / length) * 100
                    self.timeline.set(timeline_position)

                    self.current_time.config(text=self.format_ms(current))
                    self.total_time.config(text=self.format_ms(length))
            except Exception:
                pass

        self.root.after(500, self.update_timeline)

    def format_ms(self, milliseconds):
        total_seconds = int(milliseconds / 1000)
        minutes = total_seconds // 60
        seconds = total_seconds % 60
        return f"{minutes:02}:{seconds:02}"

    def start_auto_draft(self):
        if self.is_transcribing:
            messagebox.showinfo("Busy", "Auto Draft is already running.")
            return

        if not self.selected_file:
            messagebox.showerror("No file", "Please select a file first.")
            return

        if whisper is None:
            messagebox.showerror(
                "Whisper missing",
                "Whisper is not installed in this Python environment.\n\nInstall it before using Auto Draft."
            )
            return

        self.is_transcribing = True
        self.start_button.config(state="disabled", text="Drafting...")
        self.set_progress(5, "Starting Auto Draft...")
        threading.Thread(target=self.run_single_speaker_transcription, daemon=True).start()

    def run_single_speaker_transcription(self):
        try:
            self.safe_ui(lambda: self.set_progress(15, "Loading Whisper model..."))

            if self.whisper_model is None:
                self.whisper_model = whisper.load_model("base")

            self.safe_ui(lambda: self.set_progress(35, "Transcribing single-speaker dictation..."))

            result = self.whisper_model.transcribe(
                self.selected_file,
                fp16=False,
                language=None,
                task="transcribe"
            )

            text = result.get("text", "").strip()

            if not text:
                text = ""

            formatted_text = self.clean_transcript_text(text)

            self.safe_ui(lambda: self.finish_transcription(formatted_text))

        except Exception as error:
            self.safe_ui(lambda error=error: self.fail_transcription(error))

    def clean_transcript_text(self, text):
        text = text.strip()

        if not text:
            return ""

        # Keep it simple for dictation: one clean editable draft.
        return text + "\n"

    def finish_transcription(self, text):
        self.transcript_text.delete("1.0", "end")

        if text:
            self.transcript_text.insert("end", text)
        else:
            self.transcript_text.insert("end", "No speech detected. You can type manually here.")

        self.update_word_count()
        self.set_progress(100, "Auto Draft completed")
        self.set_status("Draft ready")
        self.is_transcribing = False
        self.start_button.config(state="normal", text="Start Auto Draft")
        messagebox.showinfo("Completed", "Auto Draft completed. Please proofread the text.")

    def fail_transcription(self, error):
        self.set_progress(0, "Auto Draft failed")
        self.set_status("Error")
        self.is_transcribing = False
        self.start_button.config(state="normal", text="Start Auto Draft")
        messagebox.showerror("Auto Draft Error", str(error))

    def safe_ui(self, callback):
        self.root.after(0, callback)

    def set_progress(self, value, message):
        self.progress["value"] = value
        self.progress_percent.config(text=f"{value}%")
        self.progress_title.config(text=message)

    def set_status(self, message):
        self.top_status.config(text=message)

    def update_word_count(self, event=None):
        text = self.transcript_text.get("1.0", "end-1c")
        words = len(text.split())
        self.word_count_label.config(text=f"Words: {words}")

    def get_document_text(self):
        text = self.transcript_text.get("1.0", "end-1c").strip()

        if not text or text == "Your typed document or auto draft will appear here...":
            return ""

        return text

    def copy_text(self):
        text = self.get_document_text()

        if not text:
            messagebox.showerror("No text", "There is no text to copy.")
            return

        self.root.clipboard_clear()
        self.root.clipboard_append(text)
        messagebox.showinfo("Copied", "Text copied to clipboard.")

    def clear_text(self):
        confirm = messagebox.askyesno("Clear text", "Clear the typed document area?")
        if not confirm:
            return

        self.transcript_text.delete("1.0", "end")
        self.update_word_count()

    def export_txt(self):
        text = self.get_document_text()

        if not text:
            messagebox.showerror("No text", "There is no text to export.")
            return

        save_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Text File", "*.txt")],
            title="Save Typed Document As",
            initialfile=self.default_export_name("txt")
        )

        if save_path:
            with open(save_path, "w", encoding="utf-8") as file:
                file.write(text)

            messagebox.showinfo("Success", "TXT exported successfully.")

    def export_docx(self):
        text = self.get_document_text()

        if not text:
            messagebox.showerror("No text", "There is no text to export.")
            return

        try:
            from docx import Document
        except Exception:
            messagebox.showerror("Missing package", "python-docx is not installed.")
            return

        save_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            title="Save Typed Document As",
            initialfile=self.default_export_name("docx")
        )

        if save_path:
            document = Document()
            document.add_heading("Typed Document", level=1)

            for paragraph in text.splitlines():
                document.add_paragraph(paragraph)

            document.save(save_path)
            messagebox.showinfo("Success", "DOCX exported successfully.")

    def default_export_name(self, extension):
        if self.selected_file:
            base_name = os.path.splitext(os.path.basename(self.selected_file))[0]
        else:
            base_name = "typed_document"

        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        return f"{base_name}_{timestamp}.{extension}"

    def start_fresh(self):
        if self.player:
            self.player.stop()

        self.selected_file = ""
        self.transcript_text.delete("1.0", "end")
        self.transcript_text.insert("end", "Your typed document or auto draft will appear here...")

        self.file_name_value.config(text="Name: No file selected")
        self.file_size_value.config(text="Size: -- MB")
        self.file_format_value.config(text="Format: ---")

        self.timeline.set(0)
        self.current_time.config(text="00:00")
        self.total_time.config(text="00:00")
        self.set_progress(0, "Ready")
        self.set_status("Ready")
        self.update_word_count()

    def show_history(self):
        if not self.recent_files:
            messagebox.showinfo("History", "No recent files yet.")
            return

        history_window = tk.Toplevel(self.root)
        history_window.title("Recent Files")
        history_window.geometry("560x380")
        history_window.configure(bg=COLOR_PANEL)

        tk.Label(
            history_window,
            text="Recent Files",
            font=("Segoe UI Semibold", 16),
            fg=COLOR_BLUE_2,
            bg=COLOR_PANEL
        ).pack(anchor="w", padx=18, pady=16)

        listbox = tk.Listbox(
            history_window,
            bg=COLOR_ENTRY,
            fg=COLOR_TEXT,
            font=("Segoe UI", 10),
            borderwidth=0
        )
        listbox.pack(fill="both", expand=True, padx=18, pady=(0, 18))

        for item in self.recent_files:
            listbox.insert("end", item)

        def open_selected(event=None):
            selected = listbox.curselection()
            if not selected:
                return

            name = listbox.get(selected[0])
            path = self.recent_files.get(name)

            if path and os.path.exists(path):
                self.selected_file = path
                self.load_media_for_playback(path)
                self.update_file_info(path)
                self.set_status("File ready")
                history_window.destroy()
            else:
                messagebox.showerror("Missing file", "This file no longer exists.")

        listbox.bind("<Double-Button-1>", open_selected)

    def focus_playback(self):
        self.playback_panel.focus_set()
        self.set_status("Playback")

    def focus_export(self):
        self.set_status("Export ready")
        messagebox.showinfo("Export", "Use the Export Document buttons at the bottom of the app.")

    def show_settings(self):
        messagebox.showinfo(
            "Settings",
            "Single-speaker dictation mode is enabled.\n\nAuto Draft uses Whisper Base.\nSpeaker detection is removed for this version."
        )

    def show_about(self):
        messagebox.showinfo(
            "About",
            f"{APP_NAME} {APP_VERSION}\n\nSingle-speaker dictation typing workspace.\nBuilt for typist review and correction."
        )


def main():
    root = tk.Tk()
    app = IAMTypingApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
